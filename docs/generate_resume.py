#!/usr/bin/env python3
"""Generate a formatted .docx resume matching the original PDF style."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# --- Styles setup ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.15

# Colors from original
HEADER_COLOR = RGBColor(0xBE, 0x43, 0x43)  # Dark red/maroon for section headers
NAME_COLOR = RGBColor(0x33, 0x33, 0x33)
LINK_COLOR = RGBColor(0xBE, 0x43, 0x43)
BULLET_CHAR = '\u2756'  # ❖

def add_bottom_border(paragraph, color='BE4343', width=6):
    """Add a colored bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_header(text):
    """Add a section header matching the original monospace style with underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.font.color.rgb = HEADER_COLOR
    run.bold = True
    add_bottom_border(p)

def add_job_header(title_role, date_range):
    """Add a job title line with role on left, dates on right."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    # Use a tab stop for right-aligned date
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    run = p.add_run(title_role)
    run.bold = True
    run.font.size = Pt(10)

    p.add_run('\t')

    date_run = p.add_run(date_range)
    date_run.font.size = Pt(10)
    date_run.italic = True

def add_job_description(text):
    """Add an italic description paragraph below a job header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)

def add_bullet(text):
    """Add a bullet point with ❖ character."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f'{BULLET_CHAR}  ')
    run.font.size = Pt(10)
    run.font.color.rgb = HEADER_COLOR
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), 'BE4343')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ============================================================
# NAME
# ============================================================
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
name_para.paragraph_format.space_after = Pt(0)
name_run = name_para.add_run('Alexander Root-Roatch')
name_run.font.name = 'Courier New'
name_run.font.size = Pt(26)
name_run.font.color.rgb = NAME_COLOR
name_run.bold = True

# ============================================================
# CONTACT INFO (right-aligned block)
# ============================================================
contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
contact_para.paragraph_format.space_before = Pt(0)
contact_para.paragraph_format.space_after = Pt(0)

run = contact_para.add_run('920-479-0338 | ')
run.font.size = Pt(10)
add_hyperlink(contact_para, 'arootroatch@gmail.com', 'mailto:arootroatch@gmail.com')

contact_para2 = doc.add_paragraph()
contact_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
contact_para2.paragraph_format.space_before = Pt(0)
contact_para2.paragraph_format.space_after = Pt(0)
add_hyperlink(contact_para2, 'arootroatch-dev.vercel.app', 'https://arootroatch-dev.vercel.app')

contact_para3 = doc.add_paragraph()
contact_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
contact_para3.paragraph_format.space_before = Pt(0)
contact_para3.paragraph_format.space_after = Pt(2)
add_hyperlink(contact_para3, 'linkedin.com/in/arootroatch', 'https://linkedin.com/in/arootroatch')
run = contact_para3.add_run(' | ')
run.font.size = Pt(10)
add_hyperlink(contact_para3, 'github.com/arootroatch', 'https://github.com/arootroatch')

# ============================================================
# SUMMARY
# ============================================================
add_section_header('Summary')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    'Software Craftsman who builds maintainable, test-driven systems across industries \u2014 '
    'from drone traffic management to payroll platforms to mobile commerce. I bring the same '
    'precision to software that I brought to 10 years of professional audio engineering: high '
    'standards, creative problem-solving, and a belief that great software is a craft. Passionate '
    'about Clojure, open-source tooling, and mentoring other developers.'
)
run.font.size = Pt(10)

# ============================================================
# TECHNICAL SKILLS
# ============================================================
add_section_header('Technical Skills')

skills = [
    ('Languages: ', 'Clojure, ClojureScript, Java, JavaScript, TypeScript, Python, Ruby, Babashka, VBA'),
    ('Frontend: ', 'Reagent, React, Next.js, Redux, Framer Motion, Tailwind CSS, Garden CSS'),
    ('Backend & Data: ', 'PostgreSQL, Datomic, DynamoDB, MongoDB, Redis, GraphQL, REST, WebSockets, Pub/Sub'),
    ('Infrastructure: ', 'AWS (EC2, EKS, S3, DynamoDB, IAM, SNS, SQS), Docker, Kubernetes, GitHub Actions CI/CD'),
    ('Practices: ', 'TDD/BDD, Agile, Pair Programming, Agentic Development, AI Integration'),
]

for label, items in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    bullet_run = p.add_run(f'{BULLET_CHAR}  ')
    bullet_run.font.size = Pt(10)
    bullet_run.font.color.rgb = HEADER_COLOR
    label_run = p.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(10)
    items_run = p.add_run(items)
    items_run.font.size = Pt(10)

# ============================================================
# WORK EXPERIENCE
# ============================================================
add_section_header('Work Experience')

# --- Clean Coders ---
add_job_header('Software Craftsman {Clean Coders, LLC} (Remote)', 'May 2024 \u2014 Present')
add_job_description(
    'Deliver production software for clients across UAS/drone traffic management, payroll, e-commerce, '
    'and online education at a boutique software consultancy \u2014 building, shipping, and maintaining '
    'full-stack applications as part of a 15-person engineering team.'
)

clean_coders_bullets = [
    'Co-authored ReMemory, an open-source state-management addition to C3Kit Bucket that increased front-end rendering performance 6x in ClojureScript Reagent applications',
    'Rewrote a mobile commerce backend in Clojure in 3 weeks \u2014 replacing 6 months of prior Go development \u2014 with full test coverage',
    'Built a webhook-driven Jira sync server (Clojure + Node.js Forge app) enabling a client\u2019s team and their corporate partners to stay in their own tools without losing sync',
    'Designed a cross-platform CLI in Babashka with a plugin architecture, custom test runner, local caching, and automatic task completion \u2014 integrated with an online learning platform and ran on Mac, Windows, and Linux',
    'Contributed to and maintained a suite of open-source Clojure libraries (C3Kit) for database abstraction, testing, schema validation, and API communication used across all client projects',
    'Upgraded C3Kit Wire to React 18 with idiomatic ClojureScript useEffect wrappers and an improved React testing simulator',
    'Added shutdown hooks to C3Kit Scaffold to prevent orphaned processes when used by AI agents',
    'Managed AWS infrastructure (EC2, EKS, S3, DynamoDB, IAM, SNS, SQS) with GitHub Actions CI/CD pipelines',
    'Authored technical blog posts on Kubernetes, ClojureScript Reagent patterns, and test-driven development published on cleancoders.com',
    'Recorded, edited, and appeared in educational and promotional video content for the company',
]
for b in clean_coders_bullets:
    add_bullet(b)

# --- Sound Roots ---
add_job_header(
    'Sound Roots Productions {Owner / Production Manager / Web Developer} (Nashville, TN)',
    'June 2019 \u2014 Feb 2025'
)

sound_roots_bullets = [
    'Built and maintained company websites, internal tools, and a dynamic feedback/job-application form spanning 4 venues across 3 cities',
    'Developed 3 VBA automations for payroll and invoicing that saved ~$2,000/year in labor and software costs',
    'Hired, trained, and managed audio and lighting technicians at 5 locations across 3 cities',
    'Installed and configured access points, mesh WiFi, HDMI-over-IP video matrices, and concert audio/lighting systems',
]
for b in sound_roots_bullets:
    add_bullet(b)

# --- VER ---
add_job_header(
    'VER {Audio Quality Control Technician} (Nashville, TN)',
    'Nov 2017 \u2014 June 2019'
)

ver_bullets = [
    'Built and configured large-scale audio systems for arena concerts and conferences',
    'Configured redundant Dante audio-over-ethernet and fiber-optic networks, wireless access points, and production LANs',
    'Rapidly progressed from new hire to go-to resource for system prep and troubleshooting',
]
for b in ver_bullets:
    add_bullet(b)

# ============================================================
# PROJECTS
# ============================================================
add_section_header('Projects')

# --- Gift Exchange Generator ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Gift Exchange Generator ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('{JavaScript, MongoDB, Serverless, Vitest}')
run.font.size = Pt(10)
p.add_run(' \u2014 ')
add_hyperlink(p, 'Live Site', 'https://giftexchangegenerator.netlify.app/')
p.add_run(' \u2014 ').font.size = Pt(10)
add_hyperlink(p, 'GitHub', 'https://github.com/arootroatch/ChristmasGiftExchange')

add_bullet('Framework-free Secret Santa app with a custom state-centered, event-driven architecture. 46,000+ users to date.')

# --- Tic-Tac-Toe ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Full-Stack Tic-Tac-Toe ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('{Clojure, Speclj, TDD, PostgreSQL}')
run.font.size = Pt(10)
p.add_run(' \u2014 ')
add_hyperlink(p, 'Live Site', 'https://tic-tac-toe-clojure.vercel.app/')
p.add_run(' \u2014 ').font.size = Pt(10)
add_hyperlink(p, 'GitHub', 'https://github.com/arootroatch/tic-tac-toe-clojure')

add_bullet('Fully-architected Tic-Tac-Toe supporting multiple databases (Postgres/EDN), interfaces (terminal/desktop/web), AI difficulties, and board sizes including 3D.')

# --- Java HTTP Server ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Java HTTP Server ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('{Java, JUnit}')
run.font.size = Pt(10)
p.add_run(' \u2014 ')
add_hyperlink(p, 'GitHub', 'https://github.com/arootroatch/http-server-java')

add_bullet('Multi-threaded HTTP server with zero external dependencies. Includes session management, static file serving, multipart uploads, and extensible routing.')

# ============================================================
# EDUCATION
# ============================================================
add_section_header('Education')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('University of Wisconsin Oshkosh')
run.bold = True
run.font.size = Pt(10)

add_bullet('Bachelor of Music \u2014 Recording Technology')
add_bullet('Bachelor of Arts \u2014 Spanish Language and Literature (studied abroad at the University of Salamanca, Spain)')

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/AlexRoot-Roatch/current-projects/nextjs-portfolio/docs/Root-Roatch-Resume.docx'
doc.save(output_path)
print(f'Resume saved to {output_path}')
